from __future__ import annotations
from typing import Tuple, List, Optional
import io, re
from ...core.settings import SETTINGS

_WS_RE = re.compile(r"[ \t]+")
def _squeeze_spaces(s: str) -> str:
    s = (s or "").replace("\xa0", " ")
    s = _WS_RE.sub(" ", s)
    return s.strip()

def _dbg(msg: str):
    try:
        S = SETTINGS.effective
        if bool(S().get("doc_debug", False)):
            print(f"[doc_ingest] {msg}")
    except Exception:
        pass

def _is_heading(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    m = re.match(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
    if not m:
        return None
    try:
        return max(1, min(6, int(m.group(1))))
    except Exception:
        return None

def _is_list_style(style_name: str) -> bool:
    return bool(style_name) and any(k in style_name.lower() for k in ("list", "bullet", "number"))

def _extract_paragraph_text(p) -> str:
    return _squeeze_spaces(p.text)

def extract_docx(data: bytes) -> Tuple[str, str]:
    from docx import Document
    S = SETTINGS.effective
    HEADING_MAX_LEVEL = int(S().get("docx_heading_max_level", 3))
    USE_MARKDOWN_HEADINGS = bool(S().get("docx_use_markdown_headings", True))
    PRESERVE_BULLETS = bool(S().get("docx_preserve_bullets", True))
    INCLUDE_TABLES = bool(S().get("docx_include_tables", True))
    INCLUDE_HEADERS_FOOTERS = bool(S().get("docx_include_headers_footers", False))
    MAX_PARA_CHARS = int(S().get("docx_para_max_chars", 0))
    DROP_EMPTY_LINES = bool(S().get("docx_drop_empty_lines", True))
    doc = Document(io.BytesIO(data))
    lines: List[str] = []
    try:
        title = (getattr(doc, "core_properties", None) or {}).title
        if title:
            lines.append(f"# {title}")
            lines.append("")
    except Exception:
        pass
    def _clip(s: str) -> str:
        if MAX_PARA_CHARS > 0 and len(s) > MAX_PARA_CHARS:
            return s[:MAX_PARA_CHARS] + "…"
        return s
    if INCLUDE_HEADERS_FOOTERS:
        try:
            for i, sec in enumerate(getattr(doc, "sections", []) or []):
                if i > 0:
                    break
                try:
                    hdr_ps = getattr(sec.header, "paragraphs", []) or []
                    hdr_text = "\n".join(_squeeze_spaces(p.text) for p in hdr_ps if _squeeze_spaces(p.text))
                    if hdr_text:
                        lines.append("## Header")
                        lines.append(_clip(hdr_text))
                        lines.append("")
                except Exception:
                    pass
                try:
                    ftr_ps = getattr(sec.footer, "paragraphs", []) or []
                    ftr_text = "\n".join(_squeeze_spaces(p.text) for p in ftr_ps if _squeeze_spaces(p.text))
                    if ftr_text:
                        lines.append("## Footer")
                        lines.append(_clip(ftr_text))
                        lines.append("")
                except Exception:
                    pass
        except Exception:
            pass
    for p in doc.paragraphs:
        txt = _extract_paragraph_text(p)
        if not txt and DROP_EMPTY_LINES:
            continue
        style_name = getattr(p.style, "name", "") or ""
        lvl = _is_heading(style_name)
        if lvl and lvl <= HEADING_MAX_LEVEL and USE_MARKDOWN_HEADINGS:
            prefix = "#" * max(1, min(6, lvl))
            lines.append(f"{prefix} {txt}".strip())
            continue
        if PRESERVE_BULLETS and _is_list_style(style_name):
            if txt:
                lines.append(f"- {_clip(txt)}")
            continue
        if txt:
            lines.append(_clip(txt))
        elif not DROP_EMPTY_LINES:
            lines.append("")
    if INCLUDE_TABLES and getattr(doc, "tables", None):
        for t_idx, tbl in enumerate(doc.tables):
            try:
                non_empty = any(_squeeze_spaces(cell.text) for row in tbl.rows for cell in row.cells)
            except Exception:
                non_empty = True
            if not non_empty:
                continue
            lines.append("")
            lines.append(f"## Table {t_idx + 1}")
            try:
                for row in tbl.rows:
                    cells = [_squeeze_spaces(c.text) for c in row.cells]
                    if any(cells):
                        lines.append(" | ".join(c for c in cells if c))
            except Exception:
                pass
    text = "\n".join(line.rstrip() for line in lines if line is not None).strip()
    return (text + "\n" if text else ""), "text/plain"

_RTF_CTRL_RE = re.compile(r"\\[a-zA-Z]+-?\d* ?")
_RTF_GROUP_RE = re.compile(r"[{}]")
_RTF_UNICODE_RE = re.compile(r"\\u(-?\d+)\??")
_RTF_HEX_RE = re.compile(r"\\'[0-9a-fA-F]{2}")
_HEX_BLOCK_RE = re.compile(r"(?:\s*[0-9A-Fa-f]{2}){120,}")

def _rtf_to_text_simple(data: bytes, *, keep_newlines: bool = True) -> str:
    try:
        s = data.decode("latin-1", errors="ignore")
    except Exception:
        s = data.decode("utf-8", errors="ignore")
    def _hex_sub(m):
        try:
            return bytes.fromhex(m.group(0)[2:]).decode("latin-1", errors="ignore")
        except Exception:
            return ""
    s = _RTF_HEX_RE.sub(_hex_sub, s)
    def _uni_sub(m):
        try:
            cp = int(m.group(1))
            if cp < 0:
                cp = 65536 + cp
            return chr(cp)
        except Exception:
            return ""
    s = _RTF_UNICODE_RE.sub(_uni_sub, s)
    s = s.replace(r"\par", "\n").replace(r"\line", "\n")
    s = _RTF_CTRL_RE.sub("", s)
    s = _RTF_GROUP_RE.sub("", s)
    s = _HEX_BLOCK_RE.sub("", s)
    s = s.replace("\r", "\n")
    s = re.sub(r"\n\s*\n\s*\n+", "\n\n", s)
    s = _squeeze_spaces(s)
    return s if keep_newlines else s.replace("\n", " ")

def _rtf_to_text_via_lib(data: bytes, *, keep_newlines: bool = True) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except Exception:
        return _rtf_to_text_simple(data, keep_newlines=keep_newlines)
    try:
        s = data.decode("latin-1", errors="ignore")
    except Exception:
        s = data.decode("utf-8", errors="ignore")
    try:
        txt = rtf_to_text(s)
    except Exception:
        txt = _rtf_to_text_simple(data, keep_newlines=keep_newlines)
    txt = _squeeze_spaces(txt)
    return txt if keep_newlines else txt.replace("\n", " ")

def _is_ole(b: bytes) -> bool:
    return len(b) >= 8 and b[:8] == b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"

def _generic_ole_text(data: bytes) -> str:
    S = SETTINGS.effective
    MIN_RUN = int(S().get("doc_ole_min_run_chars", 8))
    MAX_LINE = int(S().get("doc_ole_max_line_chars", 600))
    MIN_ALPHA_RATIO = float(S().get("doc_ole_min_alpha_ratio", 0.25))
    DROP_XMLISH = bool(S().get("doc_ole_drop_xmlish", True))
    DROP_PATHISH = bool(S().get("doc_ole_drop_pathish", True))
    DROP_SYMBOL_LINES = bool(S().get("doc_ole_drop_symbol_lines", True))
    DEDUPE_SHORT_REPEATS = bool(S().get("doc_ole_dedupe_short_repeats", True))
    XMLISH = re.compile(r"^\s*<[^>]+>", re.I)
    PATHISH = re.compile(r"[\\/].+\.(?:xml|rels|png|jpg|jpeg|gif|bmp|bin|dat)\b", re.I)
    SYMBOLLINE = re.compile(r"^[\W_]{6,}$")
    s = data.replace(b"\x00", b"")
    runs = re.findall(rb"[\t\r\n\x20-\x7E]{%d,}" % MIN_RUN, s)
    if not runs:
        return ""
    def _dec(b: bytes) -> str:
        try:
            return b.decode("cp1252", errors="ignore")
        except Exception:
            return b.decode("latin-1", errors="ignore")
    kept: List[str] = []
    for raw in runs:
        chunk = _dec(raw).replace("\r", "\n")
        for ln in re.split(r"\n+", chunk):
            t = ln.strip()
            if not t:
                continue
            if MAX_LINE > 0 and len(t) > MAX_LINE:
                t = t[:MAX_LINE].rstrip()
            t = _squeeze_spaces(t)
            letters = sum(1 for c in t if c.isalpha())
            if letters / max(1, len(t)) < MIN_ALPHA_RATIO:
                continue
            if DROP_XMLISH and XMLISH.search(t):
                continue
            if DROP_PATHISH and PATHISH.search(t):
                continue
            if DROP_SYMBOL_LINES and SYMBOLLINE.fullmatch(t):
                continue
            if DEDUPE_SHORT_REPEATS:
                t = re.sub(r"\b(\w{2,4})\1{2,}\b", r"\1\1", t)
            kept.append(t)
    out = "\n".join(kept)
    out = re.sub(r"\n\s*\n\s*\n+", "\n\n", out).strip()
    return out

def extract_doc_binary(data: bytes) -> Tuple[str, str]:
    head = (data[:64] or b"").lstrip()
    is_rtf = head.startswith(b"{\\rtf") or head.startswith(b"{\\RTF}")
    is_ole = _is_ole(data)
    _dbg(f"extract_doc_binary: bytes={len(data)} is_rtf={is_rtf} is_ole={is_ole}")
    if is_rtf:
        txt = _rtf_to_text_via_lib(data, keep_newlines=True).strip()
        return (txt + "\n" if txt else ""), "text/plain"
    if is_ole:
        txt = _generic_ole_text(data)
        return (txt + "\n" if txt else ""), "text/plain"
    try:
        txt = data.decode("utf-8", errors="ignore").strip()
    except Exception:
        txt = data.decode("latin-1", errors="ignore").strip()
    return (txt + ("\n" if txt else "")), "text/plain"
