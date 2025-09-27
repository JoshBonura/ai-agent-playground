# DOCX-only extraction (no .doc/RTF here)
from __future__ import annotations

import io
import re

from ...core.logging import get_logger
from ...core.settings import SETTINGS
from .ocr import ocr_image_bytes
from fastapi import HTTPException
from .ocr import MissingTesseractError, _install_hint

log = get_logger(__name__)

_WS_RE = re.compile(r"[ \t]+")


def _squeeze_spaces(s: str) -> str:
    s = (s or "").replace("\xa0", " ")
    s = _WS_RE.sub(" ", s)
    return s.strip()


def _is_heading(style_name: str) -> int | None:
    if not style_name:
        return None
    m = re.match(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
    if not m:
        return None
    try:
        return max(1, min(6, int(m.group(1))))
    except Exception:
        return None


def _is_list_style(style_name: str) -> bool:
    return bool(style_name) and any(k in style_name.lower() for k in ("list", "bullet", "number"))


def _extract_paragraph_text(p) -> str:
    return _squeeze_spaces(p.text)


def _docx_image_blobs(doc) -> list[bytes]:
    blobs: list[bytes] = []
    seen_rids = set()
    try:
        part = doc.part

        # inline images
        for ish in getattr(doc, "inline_shapes", []) or []:
            try:
                rId = ish._inline.graphic.graphicData.pic.blipFill.blip.embed
                if rId and rId not in seen_rids:
                    blob = part.related_parts[rId].blob
                    if blob:
                        blobs.append(blob)
                        seen_rids.add(rId)
            except Exception:
                pass

        for p in doc.paragraphs:
            for r in p.runs:
                for d in getattr(r._element, "xpath", lambda *_: [])(".//a:blip"):
                    try:
                        rId = d.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if rId and rId not in seen_rids:
                            blob = part.related_parts[rId].blob
                            if blob:
                                blobs.append(blob)
                                seen_rids.add(rId)
                    except Exception:
                        pass

    except Exception:
        pass
    return blobs


def extract_docx(data: bytes) -> tuple[str, str]:
    from docx import Document

    S = SETTINGS.effective
    HEADING_MAX_LEVEL = int(S().get("docx_heading_max_level", 3))
    USE_MARKDOWN_HEADINGS = bool(S().get("docx_use_markdown_headings", True))
    PRESERVE_BULLETS = bool(S().get("docx_preserve_bullets", True))
    INCLUDE_TABLES = bool(S().get("docx_include_tables", True))
    INCLUDE_HEADERS_FOOTERS = bool(S().get("docx_include_headers_footers", False))
    MAX_PARA_CHARS = int(S().get("docx_para_max_chars", 0))
    DROP_EMPTY_LINES = bool(S().get("docx_drop_empty_lines", True))

    doc = Document(io.BytesIO(data))
    lines: list[str] = []

    try:
        title = (getattr(doc, "core_properties", None) or {}).title
        if title:
            lines.append(f"# {title}")
            lines.append("")
    except Exception:
        pass

    def _clip(s: str) -> str:
        if MAX_PARA_CHARS > 0 and len(s) > MAX_PARA_CHARS:
            return s[:MAX_PARA_CHARS] + "…"
        return s

    if INCLUDE_HEADERS_FOOTERS:
        try:
            for i, sec in enumerate(getattr(doc, "sections", []) or []):
                if i > 0:
                    break
                try:
                    hdr_ps = getattr(sec.header, "paragraphs", []) or []
                    hdr_text = "\n".join(
                        _squeeze_spaces(p.text) for p in hdr_ps if _squeeze_spaces(p.text)
                    )
                    if hdr_text:
                        lines.append("## Header")
                        lines.append(_clip(hdr_text))
                        lines.append("")
                except Exception:
                    pass
                try:
                    ftr_ps = getattr(sec.footer, "paragraphs", []) or []
                    ftr_text = "\n".join(
                        _squeeze_spaces(p.text) for p in ftr_ps if _squeeze_spaces(p.text)
                    )
                    if ftr_text:
                        lines.append("## Footer")
                        lines.append(_clip(ftr_text))
                        lines.append("")
                except Exception:
                    pass
        except Exception:
            pass

    for p in doc.paragraphs:
        txt = _extract_paragraph_text(p)
        if not txt and DROP_EMPTY_LINES:
            continue
        style_name = getattr(p.style, "name", "") or ""
        lvl = _is_heading(style_name)
        if lvl and lvl <= HEADING_MAX_LEVEL and USE_MARKDOWN_HEADINGS:
            prefix = "#" * max(1, min(6, lvl))
            lines.append(f"{prefix} {txt}".strip())
            continue
        if PRESERVE_BULLETS and _is_list_style(style_name):
            if txt:
                lines.append(f"- {_clip(txt)}")
            continue
        if txt:
            lines.append(_clip(txt))
        elif not DROP_EMPTY_LINES:
            lines.append("")

    if INCLUDE_TABLES and getattr(doc, "tables", None):
        for t_idx, tbl in enumerate(doc.tables):
            try:
                non_empty = any(
                    _squeeze_spaces(cell.text) for row in tbl.rows for cell in row.cells
                )
            except Exception:
                non_empty = True
            if not non_empty:
                continue
            lines.append("")
            lines.append(f"## Table {t_idx + 1}")
            try:
                for row in tbl.rows:
                    cells = [_squeeze_spaces(c.text) for c in row.cells]
                    if any(cells):
                        lines.append(" | ".join(c for c in cells if c))
            except Exception:
                pass

        if bool(S().get("docx_ocr_images", False)):
            min_bytes = int(S().get("ocr_min_image_bytes", 16384))
        seen_ocr_text = set()
        for blob in _docx_image_blobs(doc):
            if len(blob) >= min_bytes:
                try:
                    t = (ocr_image_bytes(blob) or "").strip()
                except MissingTesseractError as e:
                    hint = _install_hint()
                    raise HTTPException(
                        status_code=424,
                        detail={
                            "code": "TESSERACT_MISSING",
                            "message": str(e),
                            "installUrl": hint["url"],
                            "note": hint["note"],
                        },
                    )
                if t:
                    key = t.lower()
                    if key not in seen_ocr_text:
                        lines.append(t)
                        seen_ocr_text.add(key)


    text = "\n".join(line.rstrip() for line in lines if line is not None).strip()
    return (text + "\n" if text else ""), "text/plain"
